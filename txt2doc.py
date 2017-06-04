from docx import Document
import re
import os

path = '/home/tusharsk/Desktop/game_is_on'
direct = os.listdir(path)

for i in direct:
    document = Document()
    document.add_heading(i, 0)
    myfile = open('/home/tusharsk/Desktop/game_is_on/'+i).read()
    myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile) # remove all non-XML-compatible characters
    p = document.add_paragraph(myfile)
    document.save('/home/tusharsk/Desktop/game_is_on'+i+'.docx')
