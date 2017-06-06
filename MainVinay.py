from docx import Document
from docx.shared import Inches
import PIL
import PIL.Image
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from cStringIO import StringIO
from docx import Document
import re
import os
import sys


def ImagetoDoc(filename):
    document = Document()
    name = raw_input('Enter name of the file which will save after conversion: ')+'.docx'
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture(filename, width = Inches(6.0))
    document.save(name)
def ImagetoPDF(filename):
    name = raw_input('Enter name of the file which will save after conversion: ')+'.pdf'
    im = PIL.Image.open(filename)
    newfilename = name
    PIL.Image.Image.save(im, newfilename, "PDF", resolution = 100.0)
def PDFtoDoc(filename):
    document = Document()
    name = raw_input('Enter name of the file which will save after conversion: ')+'.docx'
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    codec = 'utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    fp = file(filename, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password = ""
    maxpages = 0
    caching = True
    pagenos=set()

    for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password,caching=caching, check_extractable=True):
        interpreter.process_page(page)

    text = retstr.getvalue()
    myfile = text
    myfile = re.sub(r'[^\x00-\x7F]+|\x0c',r' ', myfile) # remove all non-XML-compatible characters
    illegal_xml_re = re.compile(u'[\x00-\x08\x0b-\x1f\x7f-\x84\x86-\x9f\ud800-\udfff\ufdd0-\ufddf\ufffe-\uffff]')
    myfile= illegal_xml_re.sub('', myfile)
    document.add_paragraph(myfile)
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture('q.jpg', width = Inches(6.0))
    document.save(name)
    fp.close()
    device.close()
    retstr.close()
def PDFToText(data):
    document = Document()
    fp = file(data, 'rb')
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    codec = 'utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    # Create a PDF interpreter object.
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    # Process each page contained in the document.

    for page in PDFPage.get_pages(fp):
        interpreter.process_page(page)
        data =  retstr.getvalue()
    name = raw_input('Enter name of the file which will save after conversion: ')+'.txt'  # Name of text file coerced with +.txt
    fi = open(name,'a')
    fi.write(data)
    fi.close()
def DocToText(filename):
    document=Document(filename)
    filename=filename.strip('.docx') #not able to remove .docx
    name = raw_input('Enter name of the file which will save after conversion: ')+'.txt'  # Name of text file coerced with +.txt
    fi = open(name,'a')
    f=open(filename+".txt","wb")
    for para in document.paragraphs:
        f.write(para.text)
        f.close()
    f = para.text
    fi.write(f)
    fi.close()
def TextToDoc(filename):
    document = Document()
    myfile = open(filename).read()
    myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile) # remove all non-XML-compatible characters
    p = document.add_paragraph(myfile)
    name = raw_input('Enter name of the file which will save after conversion: ')+'.docx'
    document.save(name)



i = 5;
i = int(i)
while(i != 0):
    print '1.PDFToText\n2.ImagetoDoc\n3.ImagetoPDF\n4.PDFtoDoc\n5.DocToText\n6.TextToDoc\n'
    n = raw_input("Enter Your Choice")
    print (n)
    n = int(n)
    if n == 1:
        PDFToText('a.pdf')
    elif n == 2:
        ImagetoDoc('q.jpg')
    elif n == 3:
        ImagetoPDF('q.jpg')
    elif n == 4:
        PDFtoDoc('a.pdf')
    elif n == 5:
        DocToText('demvo.docx')
    elif n == 6:
        TextToDoc('se.txt')

