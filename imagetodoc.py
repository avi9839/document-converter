from docx import Document
from docx.shared import Inches

document = Document()

p = document.add_paragraph()
r = p.add_run()
r.add_picture('q.jpg', width = Inches(6.0))
document.save('demo.docx')
